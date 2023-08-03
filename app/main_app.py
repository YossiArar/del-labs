import platform
import mammoth
import streamlit as st
import os
import json

import textwrap
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm, Pt

from app.read import add_float_picture

from docx import Document

from contants import General, Keys, Types
from PIL import Image

TEMPLATE_PATH = f'{os.getcwd()}/files/template_old.docx'
MAX_CHARS, PART1_LEN, LEFT_SIDE_MARGIN = 280, 268, 52
CM_DEFAULT_VALUES = [4.85, 4.85, 9.37, 0.75]
INCHES_DEFAULT_VALUES = [1.91, 1.91, 3.69, 0.3]
BOLD_FIELDS = ["Certificate number", 'Apprised Retail Price', "Currency"]
NOT_SPLIT_FIELDS = ["Date", "Additional comments"] + BOLD_FIELDS
PUT_RIGHT_SIDE = False


class App:
    def __init__(self):
        with open(f"{os.getcwd()}/app/config.json", 'r') as f:
            self.__config = json.loads(f.read())
        self.__doc = Document(TEMPLATE_PATH)
        self.__st = st
        self.__st.set_page_config(page_title=General.APP_NAME, layout="wide")
        self.__st.title(General.APP_NAME)
        self.__side_bar_fields = self.__initiate_ui_field_options(sidebar=True)
        self.__center_fields = self.__initiate_ui_field_options(sidebar=False)
        self.__certificate_update = {}

    def __set_columns_by_type(self, key, obj: dict):
        obj_type, value, desc = obj[Keys.TYPE], None, obj[Keys.DESC]
        label = key if obj_type == Types.CB else f"{key}{'' if not desc else f' - {desc}'}:"
        match obj_type:
            case Types.SELECT_BOX:
                value = self.__st.selectbox(label=label, options=obj[Keys.OPTIONS])
            case Types.FLOAT:
                value = self.__st.number_input(label=label, min_value=0.0)
            case Types.INT:
                value = self.__st.number_input(label=label, min_value=0)
            case Types.TXT:
                value = self.__st.text_input(label=label)
            case Types.DATE_INPUT:
                # value = str(self.__st.date_input(label=label))
                value = self.__st.date_input(label=label)
                day, month, year = f"0{value.day}" if value.day < 10 else value.day, f"0{value.month}" if value.month < 10 else value.day, value.year
                value = f"{day}-{month}-{year}"
            case Types.CB:
                value = self.__st.checkbox(label=label, value=False)
            case Types.IMAGE:
                value = self.__st.text_input(label).replace(
                    f"file:{'///' if 'Windows' in platform.platform() else '//'}", '')
                if value:
                    if os.path.exists(value):
                        if value:
                            image_data = {}
                            size_type, width, height, pos_x, pos_y, image_col = self.__st.columns(6)
                            with size_type:
                                size_ut = self.__st.radio(label='size_units', options=['CM', 'INCHES'])
                                image_data.setdefault('size_units', size_ut)
                                size_utv = CM_DEFAULT_VALUES if size_ut == 'CM' else INCHES_DEFAULT_VALUES
                                self.__st.info(f"{size_ut} default values: {size_utv}")
                            for col, label, val in zip([width, height, pos_x, pos_y],
                                                       ['width', 'height', 'position_x', 'position_y'],
                                                       size_utv):
                                with col:
                                    key_val = self.__st.number_input(label=f"{label}:", value=val)
                                    image_data.setdefault(label, key_val)
                            with image_col:
                                image = Image.open(value)
                                self.__st.image(image, caption='Your image selection')
                                if image_data:
                                    self.__certificate_update['image_data'] = image_data
                        else:
                            self.__st.error(f"The file path: {value} not exists.")
        if desc:
            value = f'{value} {desc}'
        self.__certificate_update[key] = value

    def __convert_docx_to_markdown(self, input_file, output_file=None):
        self.__doc = Document(input_file)
        paragraphs = [p.text for p in self.__doc.paragraphs]
        markdown_text = '\n'.join(paragraphs)
        output_file = output_file if output_file else input_file.replace('.docx', '_temp.md')
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
        return f, markdown_text

    def __initiate_ui_field_options(self, sidebar: bool):
        ui_fields = {}
        for key, values in self.__config.items():
            if values[Keys.SIDEBAR] == sidebar:
                ui_fields.setdefault(key, values)
        return ui_fields

    def __set_ui_fields(self, sidebar: bool):
        if not sidebar:
            with self.__st.form(key='CF'):
                for k, o in self.__center_fields.items():
                    self.__set_columns_by_type(key=k, obj=o)
                if self.__st.form_submit_button('Confirm certificate info'):
                    self.__st.info('Certificate items:')
                    self.__st.json(self.__certificate_update)
                    self.__st.session_state['data'] = self.__certificate_update
        else:
            # with self.__st.form(key='SF'):
            with self.__st.sidebar:
                for k, o in self.__side_bar_fields.items():
                    self.__set_columns_by_type(key=k, obj=o)

    def __create_new_template(self):
        doc_styles = self.__doc.styles['Normal']
        font = doc_styles.font
        font.name = 'Arial'
        font.size = Pt(7)
        left_side_values, ls_max_chars, counter = {}, 0, 0
        # doc = self.__doc.styles
        for pi, p in enumerate(self.__doc.paragraphs, 1):
            # p.paragraph_format.alignment = 2
            # print(p.text.split('\n'), p.text.count('\n'))
            for field, field_value in self.__certificate_update.items():
                if field_value is not None and len(str(field_value)) > 0:
                    field_key, line_field_margin = '{' + field + '}', 0
                    if field_key in p.text:
                        value_len_before, p_value_len, p_len = len(field_key), len(str(field_value)), len(p.text)
                        inline = p.runs
                        for i in range(len(inline)):
                            if field_key in inline[i].text:
                                if field in NOT_SPLIT_FIELDS:
                                    inline[i].text = inline[i].text.replace(field_key, str(field_value))
                                elif not PUT_RIGHT_SIDE and field not in NOT_SPLIT_FIELDS:
                                    margin = value_len_before - p_value_len  # if ls_max_chars == 0:
                                    if margin > 0:
                                        new_p_value = ',' * (margin + len(str(field_value))) + str(field_value)
                                        inline[i].text = inline[i].text.replace(field_key, new_p_value)
                                    elif margin < 0:
                                        delete_margin = ',' * (p_value_len - value_len_before)
                                        # new_p_value = ',' * (value_len_before - p_value_len) + str(field_value)
                                        inline[i].text = inline[i].text.replace(field_key, str(field_value))
                                        inline[i].text = inline[i].text.replace(delete_margin, '', 1)
                                    else:
                                        inline[i].text = inline[i].text.replace(field_key, str(field_value))
                                p.text = f"{inline[i].text}"
                                counter += 1
                                print(value_len_before, p_value_len, p_len, len(p.text))
                                # print(inline[i].text)
                    # if pi > 1:
                    #     new_p_txt = p.text
                    #     p.clear()
                    #     is_bold = True if field in BOLD_FIELDS and pi == 2 else False
                    #     p.add_run(new_p_txt, style='CommentsStyle').bold = is_bold
                    # break

            if pi == 1 and self.__certificate_update.get('Image Path'):  # Inches
                add_float_picture(p=p, image_path_or_stream=self.__certificate_update.get('Image Path'),
                                  width=self.__certificate_update.get('image_data').get('width'),
                                  height=self.__certificate_update.get('image_data').get('height'),
                                  pos_x=self.__certificate_update.get('image_data').get('position_x'),
                                  pos_y=self.__certificate_update.get('image_data').get('position_y'),
                                  size_units=self.__certificate_update.get('image_data').get('size_units'))
            print(p.text)
        new_left_side_values, first_row_rule = {}, None
        for pi, p in enumerate(self.__doc.paragraphs, 1):
            if left_side_values.get(pi) \
                    and left_side_values.get(pi)[0] in p.text \
                    and left_side_values.get(pi)[1] <= ls_max_chars:
                inline, p_obj = p.runs, left_side_values.get(pi)
                for i in range(len(inline)):
                    if p_obj[0] in inline[i].text:
                        point_margin = ',' * (ls_max_chars - p_obj[1])
                        txt_split = inline[i].text.split(',')
                        # point_margin = ',' * int((ls_max_chars - p_obj[1] + len(txt_split[-1])) / 2)
                        # point_margin = point_margin
                        new_p_txt = f"{txt_split[0]}{',' * (inline[i].text.count(',') - len(txt_split[-1]))}{point_margin}{txt_split[-1]}"
                        if first_row_rule is None and len(new_p_txt) % 2 > 0:
                            new_p_txt = new_p_txt.replace(',,', ',', 1)
                        if first_row_rule is not None:
                            if len(new_p_txt) > first_row_rule:
                                point_margin = ((len(new_p_txt) - first_row_rule) // 2) * ','
                                new_p_txt = new_p_txt.replace(point_margin, '', 1)
                            elif len(new_p_txt) < first_row_rule and len(new_p_txt) % 2 > 0:
                                point_margin = (len(new_p_txt) % 2) * ',,'
                                new_p_txt = new_p_txt.replace(',', point_margin, 1)
                        inline[i].text = inline[i].text.replace(inline[i].text, new_p_txt)
                        p.text = inline[i].text
                        new_left_side_values.setdefault(pi, [p.text, len(p.text)])
                        if first_row_rule is None:
                            first_row_rule = len(p.text)
                        break

        # last update for doc ph
        # print(left_side_values, counter)
        doc_path, doc_name = f"{os.getcwd()}/files/certificates/", f"{self.__certificate_update['Date']}_{self.__certificate_update['Certificate number']}.docx"
        doc_full_path = f"{doc_path}{doc_name}"
        self.__doc.save(doc_full_path)
        self.__convert_to_html(path=doc_full_path)

        # read data for download file
        with open(doc_full_path, 'rb') as d:
            data = d.read()
        download_button = self.__st.download_button(label='Download Certificate', data=data, file_name=doc_name)

    def __convert_to_html(self, path: str):
        with open(path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        html_path = path.replace('docx', 'html')
        with open(html_path, "w") as html_file:
            html_file.write(html)
            html_file.close()
        self.__st.success("Your new certificate is ready to use")
        self.__st.write("Click here [link](%s) to print or download file" % html_path)

    def run(self):
        # step 1 - set sidebar params
        self.__set_ui_fields(sidebar=True)
        # step 2 - set center params and waiting for user action
        self.__set_ui_fields(sidebar=False)

        # if self.__st.session_state.get('data') is not None and self.__st.button('Save'):
        if self.__st.session_state.get('data') is not None and self.__st.button('Build'):
            self.__create_new_template()
            self.__st.session_state['data'] = None


if __name__ == "__main__":
    App().run()
